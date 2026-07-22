"""
Universal Document Ingestion Pipeline
=====================================
Handles the heterogeneous formats that actually exist in an industrial plant:
PDFs (native + scanned), Word documents, Excel inspection registers, CSV
exports from CMMS, and plain text shift logs.

Every ingested document is normalised to the same internal representation:

    {
      "doc_id":    stable identifier,
      "filename":  original name,
      "doc_type":  classified category,
      "text":      full extracted plain text,
      "chunks":    [ {chunk_id, text, page, section}, ... ],
      "metadata":  {pages, size_kb, ingested_at, ...},
      "entities":  [ ... from ontology.extract_entities ... ]
    }

Chunking strategy: section-aware splitting with overlap. Industrial documents
are highly structured (numbered clauses, form fields, tabular records), so we
split on structural boundaries first and only fall back to fixed-window
splitting when no structure is detected. This preserves the semantic unit that
a retrieval hit needs to be useful -- a whole observation, not half of one.
"""

import os
import io
import re
import csv
import hashlib
import datetime

from .ontology import extract_entities

# ---------------------------------------------------------------------------
# DOCUMENT TYPE CLASSIFICATION
# ---------------------------------------------------------------------------
# Weak-supervision classifier: keyword signature voting. Deterministic and
# explainable, which matters for an audit trail.

DOC_TYPE_SIGNATURES = {
    "Work Order": [
        "work order", "wo no", "wo-", "job description", "planned start",
        "actual finish", "labour hours", "spares consumed", "maintenance type",
        "breakdown", "preventive", "cmms",
    ],
    "Incident Report": [
        "incident", "injury", "fatality", "loss of containment", "root cause",
        "immediate cause", "sequence of events", "corrective action",
        "reportable", "near miss", "lti", "first aid case",
    ],
    "Inspection Report": [
        "inspection", "thickness", "ndt", "ultrasonic", "radiography",
        "dye penetrant", "visual examination", "corrosion rate", "remaining life",
        "cml", "tml", "minimum required thickness", "next inspection due",
    ],
    "Standard Operating Procedure": [
        "standard operating procedure", "sop", "purpose", "scope",
        "responsibility", "step 1", "precaution", "do not", "shall ensure",
        "operating instruction", "startup procedure", "shutdown procedure",
    ],
    "Permit to Work": [
        "permit to work", "ptw", "hot work", "cold work", "confined space entry",
        "issuing authority", "receiving authority", "gas test", "validity",
        "isolation certificate", "lockout", "tagout", "loto",
    ],
    "Safety Datasheet": [
        "safety data sheet", "msds", "sds", "hazard identification",
        "first aid measures", "exposure controls", "flash point",
        "personal protective equipment", "un number",
    ],
    "Engineering Drawing": [
        "p&id", "piping and instrumentation", "isometric", "general arrangement",
        "drawing no", "revision", "scale", "sheet", "legend", "line list",
    ],
    "Audit / Compliance": [
        "audit", "non-conformance", "ncr", "capa", "observation",
        "compliance", "statutory", "finding", "auditor", "closure date",
        "clause", "corrective and preventive",
    ],
    "Shift Log": [
        "shift log", "shift handover", "night shift", "day shift", "general shift",
        "panel operator", "handed over", "taken over", "log entry", "hrs",
    ],
    "Equipment Manual": [
        "operation and maintenance manual", "oem", "manufacturer", "model no",
        "serial no", "lubrication schedule", "spare parts list", "torque",
        "commissioning", "installation", "technical specification",
    ],
}


def classify_document(text, filename=""):
    """
    Score text against each type signature; return (type, confidence).
    Confidence = winning score / total score, so it reflects how cleanly
    the document separated from its competitors.
    """
    haystack = (text[:12000] + " " + filename).lower()
    scores = {}
    for dtype, keywords in DOC_TYPE_SIGNATURES.items():
        score = 0
        for kw in keywords:
            hits = haystack.count(kw)
            if hits:
                # Diminishing returns: repeated keywords shouldn't dominate
                score += 1 + min(hits - 1, 3) * 0.3
        scores[dtype] = score

    total = sum(scores.values())
    if total == 0:
        return "General Document", 0.0

    best = max(scores, key=scores.get)
    confidence = scores[best] / total
    if scores[best] < 1.5:
        return "General Document", confidence
    return best, round(confidence, 3)


# ---------------------------------------------------------------------------
# FORMAT-SPECIFIC TEXT EXTRACTION
# ---------------------------------------------------------------------------

def _extract_pdf(file_bytes):
    """Native PDF text. Falls back gracefully if the page is image-only."""
    pages = []
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                txt = page.extract_text() or ""
                # Tables carry most of the value in inspection registers
                try:
                    for table in page.extract_tables():
                        rows = [
                            " | ".join((c or "").strip() for c in row)
                            for row in table if row
                        ]
                        if rows:
                            txt += "\n" + "\n".join(rows)
                except Exception:
                    pass
                pages.append((i + 1, txt))
    except Exception:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for i, page in enumerate(reader.pages):
                pages.append((i + 1, page.extract_text() or ""))
        except Exception as e:
            return [(1, f"[PDF extraction failed: {e}]")]

    # Scanned-document detection: near-zero text across pages
    total_chars = sum(len(t) for _, t in pages)
    if total_chars < 50 * max(len(pages), 1):
        ocr_pages = _try_ocr(file_bytes, len(pages))
        if ocr_pages:
            return ocr_pages
    return pages


def _try_ocr(file_bytes, n_pages):
    """
    Optional OCR path for scanned drawings and old inspection records.
    Requires tesseract binary; silently degrades if unavailable so the app
    never hard-fails on a scanned upload.
    """
    try:
        import pytesseract
        import pypdfium2 as pdfium
        from PIL import Image

        pdf = pdfium.PdfDocument(io.BytesIO(file_bytes))
        pages = []
        for i in range(min(len(pdf), 25)):  # cap for demo responsiveness
            bitmap = pdf[i].render(scale=2)
            img = bitmap.to_pil()
            txt = pytesseract.image_to_string(img)
            pages.append((i + 1, txt))
        return pages
    except Exception:
        return None


def _extract_docx(file_bytes):
    try:
        import docx
        d = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return [(1, "\n".join(parts))]
    except Exception as e:
        return [(1, f"[DOCX extraction failed: {e}]")]


def _extract_xlsx(file_bytes):
    """Each worksheet becomes its own 'page' so citations stay precise."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        pages = []
        for idx, ws in enumerate(wb.worksheets):
            lines = [f"### Sheet: {ws.title}"]
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
            pages.append((idx + 1, "\n".join(lines)))
        return pages or [(1, "")]
    except Exception as e:
        return [(1, f"[XLSX extraction failed: {e}]")]


def _extract_csv(file_bytes):
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        lines = [" | ".join(row) for row in reader if any(c.strip() for c in row)]
        return [(1, "\n".join(lines))]
    except Exception as e:
        return [(1, f"[CSV extraction failed: {e}]")]


def _extract_txt(file_bytes):
    return [(1, file_bytes.decode("utf-8", errors="replace"))]


EXTRACTORS = {
    ".pdf":  _extract_pdf,
    ".docx": _extract_docx,
    ".doc":  _extract_docx,
    ".xlsx": _extract_xlsx,
    ".xlsm": _extract_xlsx,
    ".csv":  _extract_csv,
    ".txt":  _extract_txt,
    ".md":   _extract_txt,
    ".log":  _extract_txt,
}

SUPPORTED_EXTENSIONS = sorted(EXTRACTORS.keys())


# ---------------------------------------------------------------------------
# SECTION-AWARE CHUNKING
# ---------------------------------------------------------------------------

# Structural boundaries common in industrial documents
SECTION_RE = re.compile(
    r"^\s*(?:"
    r"#{1,4}\s+.+"                          # markdown heading
    r"|\d{1,2}(?:\.\d{1,2}){0,3}[\.\)]?\s+[A-Z].{3,80}"  # 4.2.1 Clause Title
    r"|[A-Z][A-Z\s/&\-]{5,60}:?"            # ALL CAPS HEADING
    r"|(?:SECTION|CLAUSE|ANNEXURE|APPENDIX|PART)\s+[A-Z0-9]+.*"
    r")\s*$",
    re.M,
)

CHUNK_TARGET = 900      # characters -- tuned for TF-IDF retrieval granularity
CHUNK_OVERLAP = 150


def chunk_text(text, page=1):
    """
    Split into retrieval units. Structure-first, fixed-window fallback.
    Returns list of {text, page, section}.
    """
    if not text or not text.strip():
        return []

    matches = list(SECTION_RE.finditer(text))
    chunks = []

    if len(matches) >= 2:
        # Structural split
        bounds = [m.start() for m in matches] + [len(text)]
        for i in range(len(bounds) - 1):
            seg = text[bounds[i]:bounds[i + 1]].strip()
            if not seg:
                continue
            heading = seg.split("\n", 1)[0].strip()[:100]
            # Very long sections still need windowing
            if len(seg) > CHUNK_TARGET * 2:
                for sub in _window(seg):
                    chunks.append({"text": sub, "page": page, "section": heading})
            else:
                chunks.append({"text": seg, "page": page, "section": heading})
    else:
        for sub in _window(text):
            chunks.append({"text": sub, "page": page, "section": ""})

    return [c for c in chunks if len(c["text"].strip()) > 40]


def _window(text):
    """Fixed-size sliding window that respects sentence boundaries."""
    out = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + CHUNK_TARGET, n)
        if end < n:
            # Prefer to break at a sentence/newline boundary near the target
            window = text[start:end]
            for sep in ("\n\n", ". ", "\n", "; "):
                idx = window.rfind(sep)
                if idx > CHUNK_TARGET * 0.5:
                    end = start + idx + len(sep)
                    break
        out.append(text[start:end].strip())
        if end >= n:
            break
        start = max(end - CHUNK_OVERLAP, start + 1)
    return [o for o in out if o]


# ---------------------------------------------------------------------------
# MAIN INGESTION ENTRY POINT
# ---------------------------------------------------------------------------

def ingest_document(filename, file_bytes):
    """
    Full pipeline: extract -> classify -> chunk -> extract entities.
    Returns the normalised document dict, or None if the format is unsupported.
    """
    ext = os.path.splitext(filename)[1].lower()
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        return None

    pages = extractor(file_bytes)
    full_text = "\n\n".join(t for _, t in pages)

    doc_type, confidence = classify_document(full_text, filename)

    doc_id = hashlib.md5(
        (filename + str(len(file_bytes))).encode("utf-8")
    ).hexdigest()[:12]

    chunks = []
    for page_no, page_text in pages:
        for c in chunk_text(page_text, page=page_no):
            c["chunk_id"] = f"{doc_id}::c{len(chunks):04d}"
            c["doc_id"] = doc_id
            chunks.append(c)

    entities = extract_entities(full_text)

    return {
        "doc_id": doc_id,
        "filename": filename,
        "doc_type": doc_type,
        "type_confidence": confidence,
        "text": full_text,
        "chunks": chunks,
        "entities": entities,
        "metadata": {
            "pages": len(pages),
            "size_kb": round(len(file_bytes) / 1024, 1),
            "chars": len(full_text),
            "n_chunks": len(chunks),
            "n_entities": len(entities),
            "ingested_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "extension": ext,
        },
    }


def ingest_path(path):
    """Convenience loader for the bundled demo corpus."""
    with open(path, "rb") as f:
        return ingest_document(os.path.basename(path), f.read())
